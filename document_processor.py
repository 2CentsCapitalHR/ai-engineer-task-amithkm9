#!/usr/bin/env python3
"""
Document Processor Module - Enhanced Version with Improved Classification
Handles document loading, parsing, manipulation, and inline commenting
"""

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from typing import List, Dict, Optional, Tuple
import logging
from pathlib import Path
from datetime import datetime

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process and analyze ADGM legal documents with enhanced commenting"""
    
    def __init__(self):
        self.document = None
        self.document_type = None
        self.document_path = None
        self.issues = []
        self.comments_added = []
        
    def load_document(self, file_path: str) -> bool:
        """Load a Word document for processing"""
        try:
            self.document = Document(file_path)
            self.document_path = file_path
            self.document_type = self._identify_document_type()
            logger.info(f"Loaded document: {file_path}")
            logger.info(f"Identified type: {self.document_type}")
            return True
        except Exception as e:
            logger.error(f"Error loading document: {e}")
            return False
    
    def _identify_document_type(self) -> str:
        """Enhanced document type identification with improved logic"""
        if not self.document:
            return "unknown"
        
        # Get text for analysis
        text = self.get_document_text().lower()
        
        # Also check the first few paragraphs for document title
        first_paragraphs = []
        for i, para in enumerate(self.document.paragraphs[:15]):  # Check first 15 paragraphs
            if para.text.strip():
                first_paragraphs.append(para.text.lower().strip())
        
        first_text = "\n".join(first_paragraphs)
        
        # More specific pattern matching with priority order
        
        # 1. Articles of Association (HIGHEST PRIORITY)
        if any(phrase in text for phrase in [
            "articles of association",
            "article 1:",
            "article i:",
            "article 1 interpretation",
            "article 2: registered office"
        ]) or any(phrase in first_text for phrase in [
            "articles of association",
            "article 1:",
            "article i:"
        ]):
            # Additional validation for articles
            if any(indicator in text for indicator in [
                "company name",
                "registered office", 
                "share capital",
                "directors",
                "governing law",
                "interpretation"
            ]):
                return "articles_of_association"
        
        # 2. Board Resolution 
        if any(phrase in text for phrase in [
            "board resolution",
            "resolution of the board",
            "board of directors",
            "directors present",
            "it was resolved",
            "be it resolved"
        ]) and any(indicator in text for indicator in [
            "meeting",
            "directors",
            "resolved",
            "quorum"
        ]):
            return "board_resolution"
        
        # 3. Shareholder Resolution
        if any(phrase in text for phrase in [
            "shareholder resolution",
            "resolution of shareholders",
            "shareholders resolution",
            "resolution of incorporating shareholders",
            "incorporating shareholders"
        ]) or (
            "shareholder" in text and 
            "resolution" in text and
            any(indicator in text for indicator in ["shares", "shareholding", "shareholders present"])
        ):
            return "shareholder_resolution"
        
        # 4. Incorporation Application
        if any(phrase in text for phrase in [
            "adgm registration authority",
            "application for incorporation",
            "incorporation application",
            "application to incorporate",
            "company incorporation application",
            "registration authority",
            "name reservation number"
        ]) and "application" in text:
            return "incorporation_application"
        
        # 5. Employment Contract
        if any(phrase in text for phrase in [
            "employment agreement",
            "employment contract",
            "contract of employment"
        ]) or (
            "employment" in text and 
            any(term in text for term in ["employee", "employer", "salary", "working hours"])
        ):
            return "employment_contract"
        
        # 6. Register of Members and Directors
        if any(phrase in text for phrase in [
            "register of members",
            "register of directors", 
            "members register",
            "directors register",
            "part a: register",
            "part b: register"
        ]):
            return "register"
        
        # 7. UBO Declaration
        if any(phrase in text for phrase in [
            "ubo declaration",
            "beneficial ownership",
            "ultimate beneficial owner",
            "declaration of beneficial ownership"
        ]):
            return "ubo_declaration"
        
        # 8. Memorandum of Association
        if any(phrase in text for phrase in [
            "memorandum of association",
            "memorandum and articles"
        ]) or (
            "memorandum" in first_text and 
            any(term in text for term in ["name", "registered office", "objects", "liability", "share capital", "subscribers"])
        ):
            return "memorandum"
        
        # 9. General Contract/Agreement
        if any(phrase in text for phrase in [
            "this agreement",
            "this contract",
            "between party a",
            "between party b"
        ]) and any(term in text for term in ["terms and conditions", "governing law"]):
            return "commercial_agreement"
        
        # Default to general document if no specific type identified
        return "general_document"
    
    def get_document_text(self) -> str:
        """Extract all text from the document"""
        if not self.document:
            return ""
        
        full_text = []
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract text from tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return "\n".join(full_text)
    
    def add_comment_to_paragraph(self, paragraph, comment_text: str, author: str = "ADGM Corporate Agent"):
        """Add a comment to a paragraph (creates a highlighted annotation)"""
        # Since python-docx doesn't support true Word comments,
        # we'll add inline annotations with highlighting
        comment_run = paragraph.add_run(f" [COMMENT: {comment_text}]")
        comment_run.font.color.rgb = RGBColor(255, 0, 0)
        comment_run.font.italic = True
        comment_run.font.size = Pt(9)
        
        # Track that we added this comment
        self.comments_added.append({
            "text": paragraph.text[:50] + "...",
            "comment": comment_text
        })
    
    def check_and_comment_jurisdiction(self) -> List[Dict]:
        """Check for jurisdiction issues and add inline comments"""
        issues = []
        
        # Red flag jurisdictions
        incorrect_jurisdictions = {
            "UAE Federal Courts": "Per ADGM Companies Regulations 2020, Art. 6: Replace with 'ADGM Courts'",
            "Dubai Courts": "Per ADGM Companies Regulations 2020, Art. 6: Use 'ADGM Courts' instead",
            "Abu Dhabi Courts": "Per ADGM Companies Regulations 2020, Art. 6: Should be 'ADGM Courts'",
            "DIFC": "Incorrect jurisdiction - must specify 'Abu Dhabi Global Market (ADGM)'",
            "Dubai International Financial Centre": "Wrong jurisdiction - use 'Abu Dhabi Global Market'",
            "mainland UAE": "Specify 'Abu Dhabi Global Market' for ADGM entities",
            "onshore UAE": "ADGM entities must reference 'Abu Dhabi Global Market'"
        }
        
        for i, paragraph in enumerate(self.document.paragraphs):
            para_text = paragraph.text
            para_text_lower = para_text.lower()
            
            for jurisdiction, comment in incorrect_jurisdictions.items():
                if jurisdiction.lower() in para_text_lower:
                    # Special case: Allow "United Arab Emirates" when part of a full ADGM address
                    if jurisdiction == "United Arab Emirates":
                        # Check if this is part of a valid ADGM address
                        if any(adgm_ref in para_text_lower for adgm_ref in [
                            "abu dhabi global market", "adgm", "al maryah island"
                        ]):
                            continue  # Skip this as it's a valid ADGM address
                    
                    # Highlight the problematic text
                    for run in paragraph.runs:
                        if jurisdiction.lower() in run.text.lower():
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    # Add comment
                    self.add_comment_to_paragraph(paragraph, comment)
                    
                    issues.append({
                        "paragraph": i,
                        "issue": f"Incorrect jurisdiction reference: '{jurisdiction}'",
                        "severity": "high",
                        "suggestion": comment,
                        "regulation": "ADGM Companies Regulations 2020, Art. 6",
                        "source": "Rule-based Check"
                    })
        
        # Check if ADGM is mentioned at all (skip for general documents)
        text = self.get_document_text().lower()
        adgm_patterns = ["abu dhabi global market", "adgm"]
        has_adgm = any(pattern in text for pattern in adgm_patterns)
        
        # Only flag missing ADGM for specific document types that require it
        adgm_required_types = [
            "articles_of_association", 
            "board_resolution", 
            "shareholder_resolution",
            "memorandum",
            "incorporation_application",
            "employment_contract"
        ]
        
        if not has_adgm and self.document_type in adgm_required_types:
            # Add comment to first paragraph
            if self.document.paragraphs:
                self.add_comment_to_paragraph(
                    self.document.paragraphs[0],
                    "Missing ADGM jurisdiction - Per ADGM Companies Regulations 2020, Art. 6: Must specify 'Abu Dhabi Global Market'"
                )
            
            issues.append({
                "paragraph": 0,
                "issue": "Missing ADGM jurisdiction reference",
                "severity": "high",
                "suggestion": "Add explicit reference to 'Abu Dhabi Global Market (ADGM)' jurisdiction",
                "regulation": "ADGM Companies Regulations 2020, Art. 6",
                "source": "Rule-based Check"
            })
        
        return issues
    
    def check_and_comment_weak_language(self) -> List[Dict]:
        """Check for weak language and add inline comments"""
        issues = []
        
        weak_terms = {
            "may": ("shall", "Per ADGM legal drafting standards: Use 'shall' for mandatory obligations"),
            "might": ("shall", "Per ADGM legal drafting standards: Replace with 'shall' for binding effect"),
            "could": ("shall", "Per ADGM legal drafting standards: Use 'shall' for mandatory provisions"),
            "possibly": ("shall", "Ambiguous language - use 'shall' for clarity"),
            "perhaps": ("shall", "Uncertain language - replace with 'shall'"),
            "should": ("shall", "Weak obligation - use 'shall' for binding requirements")
        }
        
        # Skip weak language check for certain phrases where these terms are acceptable
        acceptable_contexts = [
            "may be called",  # Common in meeting provisions
            "as may be",      # Common legal phrasing
            "may from time to time",  # Standard legal language
            "shall have the power",  # When followed by strong language
            "may terminate",  # Standard in termination clauses
            "may be amended"  # Standard in amendment clauses
        ]
        
        for i, paragraph in enumerate(self.document.paragraphs):
            para_text = paragraph.text
            para_lower = para_text.lower()
            
            # Skip if paragraph contains acceptable contexts
            if any(context in para_lower for context in acceptable_contexts):
                continue
            
            for weak, (strong, comment) in weak_terms.items():
                # Use word boundaries to avoid false positives
                pattern = r'\b' + re.escape(weak) + r'\b'
                if re.search(pattern, para_text, re.IGNORECASE):
                    # Highlight weak terms
                    for run in paragraph.runs:
                        if re.search(pattern, run.text, re.IGNORECASE):
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    
                    # Add comment
                    self.add_comment_to_paragraph(paragraph, comment)
                    
                    issues.append({
                        "paragraph": i,
                        "issue": f"Weak language detected: '{weak}'",
                        "severity": "medium",
                        "suggestion": f"Replace '{weak}' with '{strong}'",
                        "context": para_text[:100],
                        "regulation": "ADGM legal drafting standards",
                        "source": "Rule-based Check"
                    })
        
        return issues
    
    def check_and_comment_required_sections(self) -> List[Dict]:
        """Check for required sections and add comments for missing ones"""
        issues = []
        text = self.get_document_text().lower()
        
        # Updated required sections with correct requirements for each document type
        required_sections = {
            "articles_of_association": {
                "company name": "Per ADGM Companies Regulations 2020, Art. 30: Company name required",
                "registered office": "Per ADGM Companies Regulations 2020, Art. 25: Registered office must be specified",
                "share capital": "Per ADGM Companies Regulations 2020, Art. 12: Share capital details required",
                "directors": "Per ADGM Companies Regulations 2020, Art. 15: Director provisions required",
                "governing law": "Per ADGM Companies Regulations 2020, Art. 6: Governing law clause required",
                "interpretation": "Definitions section required for clarity"
            },
            "board_resolution": {
                "date": "Date of resolution required",
                "present": "Attendance record required",
                "resolved": "Resolution language required",
                "signature": "Director signatures required"
            },
            "shareholder_resolution": {
                "shareholder": "Shareholder details required",
                "resolved": "Resolution language required",
                "signature": "Shareholder signatures required"
            },
            "memorandum": {
                "name": "Company name required",
                "registered office": "Registered office required",
                "objects": "Objects of the company required",
                "liability": "Liability of members required",
                "share capital": "Share capital required",
                "subscriber": "Subscriber details required"
            },
            "incorporation_application": {
                "company details": "Company information section required",
                "registered office": "ADGM registered office address required",
                "share capital": "Share capital structure required",
                "directors": "Director information required",
                "shareholders": "Shareholder details required"
            },
            "employment_contract": {
                "employee": "Employee details required",
                "position": "Job position/title required",
                "salary": "Compensation details required",
                "working hours": "Working hours specification required",
                "termination": "Termination provisions required"
            }
        }
        
        if self.document_type in required_sections:
            missing_sections = []
            
            for section, comment in required_sections[self.document_type].items():
                # Use more flexible matching for sections
                section_found = False
                
                # Check for various forms of the section
                if section in text or section.replace(" ", "") in text:
                    section_found = True
                
                # Special checks for common variations
                if not section_found:
                    if section == "date" and any(d in text for d in ["dated", "date:", "on this day", "january", "february", "march", "april", "may", "june", "july", "august", "september", "october", "november", "december", "2024", "2025"]):
                        section_found = True
                    elif section == "present" and any(p in text for p in ["present:", "attendance", "directors present", "in attendance"]):
                        section_found = True
                    elif section == "resolved" and any(r in text for r in ["resolved", "resolution", "it was resolved", "be it resolved"]):
                        section_found = True
                    elif section == "signature" and any(s in text for s in ["signature", "signed", "____", "authorized signatory", "signatory"]):
                        section_found = True
                    elif section == "shareholder" and any(s in text for s in ["shareholder", "member", "shares", "shareholding"]):
                        section_found = True
                    elif section == "employee" and any(e in text for e in ["employee", "employment", "employer"]):
                        section_found = True
                    elif section == "salary" and any(s in text for s in ["salary", "compensation", "remuneration", "aed", "usd"]):
                        section_found = True
                    elif section == "company name":
                        # Enhanced company name detection
                        # Check document title, headers, and first few paragraphs
                        first_paragraphs_text = " ".join([p.text for p in self.document.paragraphs[:5]]).lower()
                        if any(indicator in first_paragraphs_text for indicator in [
                            "limited", "ltd", "llc", "inc", "corporation", "corp", "company"
                        ]) or any(name_pattern in text for name_pattern in [
                            "company name", "company:", "entity name", "\"company\" means"
                        ]):
                            section_found = True
                
                if not section_found:
                    missing_sections.append((section, comment))
                    
                    issues.append({
                        "paragraph": -1,
                        "issue": f"Missing required section: '{section}'",
                        "severity": "high",
                        "suggestion": f"Add a section covering '{section}'",
                        "regulation": comment,
                        "source": "Rule-based Check"
                    })
            
            # Add a summary comment at the beginning if sections are missing
            if missing_sections and self.document.paragraphs:
                summary_comment = f"Missing {len(missing_sections)} required sections: " + \
                                 ", ".join([s[0] for s in missing_sections[:3]])
                if len(missing_sections) > 3:
                    summary_comment += f" and {len(missing_sections) - 3} more"
                
                self.add_comment_to_paragraph(
                    self.document.paragraphs[0],
                    summary_comment
                )
        
        return issues
    
    def check_and_comment_signatory_sections(self) -> List[Dict]:
        """Check signatory sections and add comments"""
        issues = []
        text = self.get_document_text().lower()
        
        # Skip signature check for documents that might not require it
        skip_signature_check = ["general_document", "register"]
        if self.document_type in skip_signature_check:
            return issues
        
        # Check for signature blocks
        signature_indicators = ["signature", "signed", "authorized signatory", "_______", "____"]
        has_signature_section = any(indicator in text for indicator in signature_indicators)
        
        if not has_signature_section:
            # Add comment to last paragraph
            if self.document.paragraphs:
                self.add_comment_to_paragraph(
                    self.document.paragraphs[-1],
                    "Per ADGM execution requirements: Add signature blocks with name, title, and date fields"
                )
            
            issues.append({
                "paragraph": len(self.document.paragraphs) - 1,
                "issue": "Missing signature section",
                "severity": "high",
                "suggestion": "Add proper signature blocks with name, title, and date fields",
                "regulation": "ADGM execution requirements",
                "source": "Rule-based Check"
            })
        else:
            # Check for incomplete signatures (look for empty signature lines)
            for i, paragraph in enumerate(self.document.paragraphs):
                para_text = paragraph.text
                
                # Check if this looks like an incomplete signature block
                if ("_______" in para_text or "____" in para_text) and "name:" not in para_text.lower():
                    # Check if the signature block is missing name/title/date
                    signature_lines = []
                    for j in range(max(0, i-2), min(len(self.document.paragraphs), i+3)):
                        signature_lines.append(self.document.paragraphs[j].text.lower())
                    
                    signature_text = " ".join(signature_lines)
                    
                    missing_elements = []
                    if "name:" not in signature_text and not any(name in signature_text for name in ["john", "sarah", "michael", "emma", "omar", "alice", "bob", "david", "maria", "jennifer", "robert", "alexandra", "chen"]):
                        missing_elements.append("signatory name")
                    if "date:" not in signature_text:
                        missing_elements.append("date field")
                    
                    if missing_elements:
                        self.add_comment_to_paragraph(
                            paragraph,
                            f"Incomplete signature block - missing: {', '.join(missing_elements)}"
                        )
                        
                        issues.append({
                            "paragraph": i,
                            "issue": f"Incomplete signature block - missing {', '.join(missing_elements)}",
                            "severity": "medium",
                            "suggestion": "Complete all signature fields with name, title, and date",
                            "regulation": "ADGM documentation standards",
                            "source": "Rule-based Check"
                        })
        
        return issues
    
    def perform_comprehensive_review(self) -> List[Dict]:
        """Perform all checks and add inline comments"""
        all_issues = []
        
        # Log the document type for debugging
        logger.info(f"Performing review for document type: {self.document_type}")
        
        # Run all checks with commenting
        all_issues.extend(self.check_and_comment_jurisdiction())
        all_issues.extend(self.check_and_comment_weak_language())
        all_issues.extend(self.check_and_comment_required_sections())
        all_issues.extend(self.check_and_comment_signatory_sections())
        
        # Store issues for later reference
        self.issues = all_issues
        
        logger.info(f"Review complete. Found {len(all_issues)} issues, added {len(self.comments_added)} comments")
        
        return all_issues
    
    def save_reviewed_document(self, output_path: str) -> bool:
        """Save the reviewed document with all comments and highlights"""
        if not self.document:
            return False
        
        try:
            # Create a new document for the reviewed version
            reviewed_doc = Document()
            
            # Add review header
            header = reviewed_doc.add_heading('ADGM COMPLIANCE REVIEW REPORT', 0)
            header.alignment = 1  # Center
            
            # Add metadata
            reviewed_doc.add_paragraph(f"Review Date: {datetime.now().strftime('%B %d, %Y')}")
            reviewed_doc.add_paragraph(f"Document Type: {self.document_type.replace('_', ' ').title()}")
            reviewed_doc.add_paragraph(f"Total Issues Found: {len(self.issues)}")
            
            # Add severity breakdown
            high_issues = sum(1 for i in self.issues if i.get("severity") == "high")
            medium_issues = sum(1 for i in self.issues if i.get("severity") == "medium")
            low_issues = sum(1 for i in self.issues if i.get("severity") == "low")
            
            reviewed_doc.add_paragraph(f"High Severity: {high_issues}")
            reviewed_doc.add_paragraph(f"Medium Severity: {medium_issues}")
            reviewed_doc.add_paragraph(f"Low Severity: {low_issues}")
            
            # Add separator
            reviewed_doc.add_paragraph("=" * 70)
            reviewed_doc.add_heading('REVIEWED DOCUMENT WITH INLINE COMMENTS', 1)
            reviewed_doc.add_paragraph("=" * 70)
            
            # Copy the original document with comments
            for paragraph in self.document.paragraphs:
                reviewed_doc.add_paragraph(paragraph.text)
            
            # Add summary of comments at the end
            if self.comments_added:
                reviewed_doc.add_page_break()
                reviewed_doc.add_heading('COMMENT SUMMARY', 1)
                
                for i, comment_info in enumerate(self.comments_added, 1):
                    reviewed_doc.add_paragraph(f"{i}. Location: {comment_info['text']}")
                    comment_para = reviewed_doc.add_paragraph(f"   Comment: {comment_info['comment']}")
                    comment_para.paragraph_format.left_indent = Inches(0.5)
                    reviewed_doc.add_paragraph("")
            
            # Save the document
            reviewed_doc.save(output_path)
            logger.info(f"Saved reviewed document to: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error saving document: {e}")
            return False